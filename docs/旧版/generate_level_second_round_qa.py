from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

out = Path(r'c:/Users/loofnn/Desktop/sdc/策划案文档/关卡地图说明书/关卡说明书第二轮问答.docx')


def set_font(run, size=10.5, bold=False, color=None):
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def setup(doc):
    styles = doc.styles
    styles['Normal'].font.name = '宋体'
    styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    styles['Normal'].font.size = Pt(10.5)
    for name, size, color in [('Heading 1', 16, '1F4E79'), ('Heading 2', 13, '1F4E79'), ('Heading 3', 11, '000000')]:
        st = styles[name]
        st.font.name = '宋体'
        st._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
    for sec in doc.sections:
        sec.top_margin = Cm(2.2)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.4)
        sec.right_margin = Cm(2.4)


def p(doc, text='', style=None, size=10.5, bold=False, color=None, align=None):
    para = doc.add_paragraph(style=style)
    para.paragraph_format.line_spacing = 1.25
    if align:
        para.alignment = align
    r = para.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    return para


def bullet(doc, items):
    for item in items:
        para = doc.add_paragraph(style='List Bullet')
        para.paragraph_format.line_spacing = 1.18
        r = para.add_run(item)
        set_font(r)


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run(h)
        set_font(r, bold=True)
        shade_cell(cell, 'D9EAF7')
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            para = cells[i].paragraphs[0]
            r = para.add_run(str(val))
            set_font(r, size=10)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return t


def reply_box(doc, label='你的回复 / 修改意见：'):
    p(doc, label, bold=True, color='C00000')
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    cell = t.rows[0].cells[0]
    cell.text = '\n\n\n'
    shade_cell(cell, 'FFF2CC')
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def q(doc, text):
    p(doc, text, bold=True)
    reply_box(doc)


doc = Document()
setup(doc)

p(doc, '关卡说明书第二轮问答', size=22, bold=True, color='1F4E79', align=WD_ALIGN_PARAGRAPH.CENTER)
p(doc, '用于补齐真实专项策划案所需的执行级规则、交互逻辑、程序规则、表格方向与美术资源清单', size=12, color='595959', align=WD_ALIGN_PARAGRAPH.CENTER)
p(doc, '')
table(doc, ['文档项', '内容'], [
    ['文档定位', '关卡说明书第二轮问答'],
    ['当前目标', '补齐关卡专项策划案落地所需的细节点'],
    ['输出用途', '在你回复后，重写真实项目标准版《关卡说明书》'],
    ['回复方式', '请直接在黄色区域填写回复、修改意见、示例或“待定”'],
])

doc.add_heading('一、章节入口与开放逻辑', level=1)
q(doc, '1.1 关卡系统入口在哪里？是主界面固定入口、家园入口、章节页入口，还是多个入口并存？')
q(doc, '1.2 玩家点击关卡入口后，先进入“章节列表页”还是直接进入“当前推荐章节页”？')
q(doc, '1.3 章节开放条件如何判定？例如角色等级、前置章节通关、任务、世界进度、系统开放等级。')
q(doc, '1.4 若章节未解锁，客户端如何表现？例如灰态、锁图标、提示文本、跳转来源。')
q(doc, '1.5 若章节已解锁但未通关，客户端如何表现？若已通关，是否需要额外显示通关状态、推荐难度、推荐系数分？')
q(doc, '1.6 章节页面需要展示哪些信息？例如章节名称、章节主题、推荐战力、可掉落装备池、Boss 预览、难度 Tag 入口。')


doc.add_heading('二、自选难度 Tag 页面与交互', level=1)
q(doc, '2.1 玩家进入章节前，自选难度 Tag 页面是强制弹出，还是可选展开页面？')
q(doc, '2.2 Tag 页面需要展示哪些信息？例如 Tag 名称、效果说明、系数分、解锁条件、当前是否可选、互斥提示。')
q(doc, '2.3 Tag 的未解锁态、已解锁未选态、已选态、互斥不可选态，客户端分别如何表现？')
q(doc, '2.4 玩家点击一个 Tag 时，交互逻辑是什么？是否立即选中？是否需要二次确认？若与其他 Tag 冲突如何提示？')
q(doc, '2.5 页面上是否需要实时显示“当前已选 Tag 数量”“当前系数分总值”“当前奖励预览档位”？')
q(doc, '2.6 玩家点击“开始挑战”前，是否需要校验条件？例如至少选 0 个 Tag、Tag 上限、章节开放、门票、体力类资源（如果没有就写无）。')
q(doc, '2.7 如果玩家没有选择任何 Tag，是否允许直接进入？是否有默认难度提示？')


doc.add_heading('三、章节进入与开局流程', level=1)
q(doc, '3.1 玩家点击“开始挑战”后，进入关卡前的完整流程是什么？例如加载页 → 章节开场表现 → 出生房间。')
q(doc, '3.2 玩家首次进入章节时，是否需要播放章节介绍、Boss 介绍、章节主题提示或引导弹窗？')
q(doc, '3.3 开局房间是否固定？是否存在安全房、教学房、初始事件房？')
q(doc, '3.4 开局玩家默认拥有的资源有哪些？例如初始背包容量、局内货币、默认状态、初始药剂、初始装备状态。')
q(doc, '3.5 如果玩家是中途断线重连、关闭游戏后继续，进入章节时如何恢复？是回到当前房间开局、房间重置、还是保持原状态？')


doc.add_heading('四、房间推进规则', level=1)
q(doc, '4.1 房间完成条件有哪些类型？例如清怪完成、交互完成、倒计时完成、保护目标完成。当前第一版要支持哪些类型？')
q(doc, '4.2 普通战斗房的完成条件是否为“击败全部怪物”？精英房、深渊房是否也一样？')
q(doc, '4.3 房间完成后，掉落是在房间内直接拾取，还是结算后统一弹出奖励？是否存在“未拾取自动进背包”规则？')
q(doc, '4.4 若玩家背包满了，房间奖励如何处理？掉在地上、禁止拾取、弹提示，还是转邮件/结算缓存？')
q(doc, '4.5 房间中若存在事件交互物，交互逻辑如何定义？例如靠近触发、点击按钮触发、清怪后解锁触发。')
q(doc, '4.6 房间内是否允许中途退出？若允许，退出按钮在哪里，点击后是否二次确认，退出后按什么规则结算？')


doc.add_heading('五、三选一传送门逻辑', level=1)
q(doc, '5.1 当前房间完成后，三选一传送门是固定出现 3 个，还是根据配置可能出现 2 个 / 3 个？')
q(doc, '5.2 三选一传送门出现时，玩家界面如何表现？是直接在场景里立门，还是弹出独立选择界面，还是两者结合？')
q(doc, '5.3 每个传送门上需要显示哪些信息？例如房间类型、奖励类型、危险等级、推荐标签、特殊标识。')
q(doc, '5.4 玩家选择传送门后，是否立刻进入下一房间？还是先播放确认动画、确认弹窗或预览页？')
q(doc, '5.5 若玩家在传送门界面停留不选，是否允许打开背包、整理装备、查看地图、查看当前 Build？')
q(doc, '5.6 是否存在“三选一刷新”机制？例如消耗道具刷新一次当前房间选项。若当前没有则写无。')


doc.add_heading('六、特殊传送门与特殊房间', level=1)
q(doc, '6.1 特殊传送门是在常规三选一之外额外出现，还是会替换其中一个常规选项？')
q(doc, '6.2 特殊传送门出现时，客户端是否需要额外特效、额外文案、额外提示动画？')
q(doc, '6.3 玩家点击特殊传送门前，能看到哪些信息？例如“特殊事件”“深渊关卡”“隐藏奖励”“未知危险”。')
q(doc, '6.4 特殊传送门和三选一同时存在时，优先级如何定义？玩家是否可以放弃特殊门、仍然走普通三选一？')
q(doc, '6.5 特殊传送门进入后，是否还能回到原路线？如果不能，客户端是否要给出风险提示？')
q(doc, '6.6 特殊房间库目前你计划先支持哪些类型？例如深渊主题房、隐藏奖励房、角色事件房、时装事件房。')


doc.add_heading('七、每层 Boss 后流程', level=1)
q(doc, '7.1 击败 Boss 后，完整流程顺序是什么？请按顺序列出，例如：Boss 死亡动画 → 房间结算 → 翻牌 → 商店 → 传输 → 下一层。')
q(doc, '7.2 Boss 死亡后，玩家是否还能在场景内移动、拾取、开背包、查看属性？')
q(doc, '7.3 翻牌奖励界面的交互是什么？例如显示几张牌、是否只能选 1 张、选中后是否有确认、未选能否关闭。')
q(doc, '7.4 翻牌奖励被选中后，奖励是直接进局外背包、局内背包，还是先进临时奖励栏？')
q(doc, '7.5 道中商店界面包含哪些页签或模块？例如购买页、出售页、局内道具页、稀有资源页。')
q(doc, '7.6 商店里的商品刷新逻辑是什么？固定池、权重随机、受章节 / Tag / 层数影响？')
q(doc, '7.7 对外传输界面的交互是什么？例如从局内背包勾选若干道具 → 点击确认 → 成功提示。')
q(doc, '7.8 若玩家超过可传输数量，客户端如何限制？若传输后后悔，是否允许撤销？')
q(doc, '7.9 玩家完成翻牌、商店、传输后，如何进入下一层？是点击“继续挑战”按钮，还是自动进入传送门？')


doc.add_heading('八、胜利、失败、主动退出、隐藏 Boss 优先级', level=1)
q(doc, '8.1 请你按程序判定顺序描述：普通失败、Boss 胜利、隐藏 Boss 触发、主动退出，这几种结果的优先级如何排列？')
q(doc, '8.2 玩家在房间战斗中主动退出，结算时是否保留当前背包中的全部候选道具，再按失败带出数量筛选？')
q(doc, '8.3 玩家在 Boss 后商店阶段主动退出，是否也按失败结算？此时翻牌奖励和已传输道具分别如何处理？')
q(doc, '8.4 隐藏 Boss 触发后，玩家选择“不挑战”时，是否直接进入胜利结算页？')
q(doc, '8.5 隐藏 Boss 挑战失败时，是否覆盖原本的胜利结果，按失败结算？已传输奖励是否仍然安全保留？')
q(doc, '8.6 胜利结算页与失败结算页分别需要展示哪些信息？例如本局收益、挑战层数、已传输物品、可带出数量、系数分评价。')


doc.add_heading('九、关卡系统的程序规则补细', level=1)
q(doc, '9.1 房间生成规则目前你希望我写到什么程度？例如“先读章节表 → 再读层表 → 再抽房间池 → 再抽事件”。')
q(doc, '9.2 事件房与普通战斗房的抽取优先级，是否存在必出、保底、互斥、层级限制？')
q(doc, '9.3 深渊房、商人房、休息房、特殊房，哪些可以同层重复出现，哪些需要限制次数？')
q(doc, '9.4 Boss 房是固定终点还是也走表生成？若固定，Boss 本体是章节固定还是层固定？')
q(doc, '9.5 当前关卡系统需要哪些异常规则？例如：背包满、奖励未领、断线重连、传输中断、翻牌未完成关闭游戏。')
q(doc, '9.6 客户端与服务端（或逻辑层）职责是否需要在策划案里初步区分？例如哪些是表现层、哪些是判定层。')


doc.add_heading('十、关卡系统表格方向（第一版）', level=1)
p(doc, '这一部分我先不要求你给出完整字段，只想先和你确认“应有哪些表”。后续我会先输出表和字段方向版。')
q(doc, '10.1 你认同关卡系统至少需要这些表吗：章节表、层配置表、房间池表、事件权重表、特殊传送门表、Boss 奖励表、翻牌奖励表、传输规则表、Tag 表？')
q(doc, '10.2 除了上面这些，你还觉得关卡系统必须额外有哪几张表？')
q(doc, '10.3 “地图事件会关联一个关卡ID，从该关卡ID中调用怪物组成”这一规则，是否应该形成单独的‘关卡事件-关卡ID映射表’或类似结构？')
q(doc, '10.4 房间类型、奖励类型、危险等级、特殊标记这些展示信息，是放在房间池表里，还是拆成单独表现配置表？')
q(doc, '10.5 一阶段你希望我在正式版关卡说明书里，把表格方向写到什么粒度：表名 + 作用，还是表名 + 核心字段方向？')


doc.add_heading('十一、美术资源清单（关卡系统）', level=1)
q(doc, '11.1 章节系统目前至少需要哪些美术资源？例如章节入口图、章节背景图、章节图标、Boss 预览图。')
q(doc, '11.2 房间推进系统至少需要哪些资源？例如三选一传送门 UI、普通传送门、特殊传送门、危险等级标识、奖励类型图标。')
q(doc, '11.3 Boss 后流程至少需要哪些资源？例如翻牌界面、牌面图、商店界面、传输界面、结算界面。')
q(doc, '11.4 特殊房间目前至少需要哪些资源清单？例如深渊房底图、特殊门特效、隐藏房提示图标。')
q(doc, '11.5 你希望我在正式版《关卡说明书》中，美术资源清单按“界面资源 / 场景资源 / 特效资源 / icon资源”分类书写吗？')


doc.add_heading('十二、我准备如何落地下一版《关卡说明书》', level=1)
p(doc, '你回复本轮后，我会按以下模式重写正式版：')
bullet(doc, [
    '文档总览',
    '功能规则',
    '流程图说明',
    '交互逻辑拆页写法',
    '程序规则拆页写法',
    '表格方向说明',
    '美术资源清单',
])
p(doc, '这会是第一版接近真实项目使用标准的《关卡说明书》。')

doc.save(out)
print(out.as_posix())
