from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

base = Path(r'c:/Users/loofnn/Desktop/sdc/项目说明书')


def set_font(run, size=10.5, bold=False, color=None):
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def setup(doc):
    styles = doc.styles
    styles['Normal'].font.name = '宋体'
    styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    styles['Normal'].font.size = Pt(10.5)
    for name, size, color in [('Heading 1', 16, '1F4E79'), ('Heading 2', 13, '1F4E79'), ('Heading 3', 11, '000000')]:
        st = styles[name]
        st.font.name = '宋体'
        st._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
    for sec in doc.sections:
        sec.top_margin = Cm(2.2)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.4)
        sec.right_margin = Cm(2.4)


def p(doc, text='', style=None, size=10.5, bold=False, color=None, align=None):
    para = doc.add_paragraph(style=style)
    para.paragraph_format.line_spacing = 1.25
    if align:
        para.alignment = align
    r = para.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    return para


def bullet(doc, items):
    for item in items:
        para = doc.add_paragraph(style='List Bullet')
        para.paragraph_format.line_spacing = 1.18
        r = para.add_run(item)
        set_font(r)


def num(doc, items):
    for item in items:
        para = doc.add_paragraph(style='List Number')
        para.paragraph_format.line_spacing = 1.18
        r = para.add_run(item)
        set_font(r)


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run(h)
        set_font(r, bold=True)
        shade_cell(cell, 'D9EAF7')
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            para = cells[i].paragraphs[0]
            r = para.add_run(str(val))
            set_font(r, size=10)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return t


def cover(doc, title, subtitle):
    p(doc, title, size=22, bold=True, color='1F4E79', align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, subtitle, size=12, color='595959', align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, '')

# 1. 项目记忆与工作上下文文档
mem = Document()
setup(mem)
cover(mem, 'SDC项目记忆与工作上下文', '用于沉淀项目关键记忆、协作习惯、可用技能和阅读入口')
table(mem, ['文档项', '内容'], [
    ['文档定位', '项目记忆 / 工作上下文 / 阅读索引'],
    ['主要用途', '帮助快速理解项目背景、协作习惯、关键变更与文档入口'],
    ['更新原则', '仅保留稳定信息；具体数值、系统细节以后续专项策划案和表格为准'],
])
mem.add_heading('1. 当前项目关键记忆', level=1)
bullet(mem, [
    '项目类型：像素俯视角、暗黑-like 刷宝、肉鸽 ARPG。',
    '产品体量：中体量买断制；首发平台：PC / Steam。',
    '核心支柱：装备刷宝、装备驱动技能裂变、随机地图结构、自选难度 Tag、永久养成保留。',
    '地图推进：采用哈迪斯-like 三选一路线推进，不可返回，迷雾机制已取消。',
    '旧版疲劳值系统已取消，由自选难度 Tag、事件权重、地图推进、背包与结算机制承接压力与收益。',
    '层结算机制已改为：翻牌奖励 + 道中商店 + 对外传输 + 全局胜败带出。',
    '完整 GDD 写作原则：方向清楚 + 关键规则 + 详细数值留给配置表 / 专项策划案。',
])
mem.add_heading('2. 当前协作习惯', level=1)
bullet(mem, [
    '长文本、大量问题优先使用 Word 作为中间介质；Markdown 作为备选。',
    '先讨论并确认框架，再撰写正式文档。',
    '自动推进常规工作，只有在阻碍、歧义或高风险操作时再确认。',
    '项目说明书、GDD、专项策划案分别承担不同层级职责，不混写。',
])
mem.add_heading('3. 当前可直接阅读的核心文档', level=1)
table(mem, ['文档名称', '用途'], [
    ['SDC项目说明书.docx', '项目整体框架说明与系统边界声明。'],
    ['SDC完整GDD问题与框架确认稿.docx', '用于收集 GDD 两轮问题回复与结构对齐。'],
    ['SDC完整GDD初稿.docx', '当前完整 GDD 初稿。'],
    ['SDC项目开发Plan.docx', '基于 GDD 生成的开发阶段计划与模块优先级。'],
    ['SDC第一版Demo说明书.docx', '以关卡内循环验证为主的第一版 Demo 说明书。'],
])
mem.add_heading('4. 当前建议拆分的专项文档', level=1)
table(mem, ['专项文档', '主要内容'], [
    ['战斗规则说明', '战斗公式、属性、表现、技能整体架构。'],
    ['技能说明', '具体技能机制、裂变方式、技能表格匹配。'],
    ['装备系统策划案', '装备字段、品质、词条、套装、洗练、提取、镶嵌、融合。'],
    ['关卡策划案', '章节、房间、路线选择、事件权重、特殊传送门。'],
    ['玩家角色策划案', '职业、转职、觉醒、天赋、技能槽位。'],
    ['怪物策划案', '普通怪、精英怪、深渊、Boss、掉落。'],
    ['家园系统策划案', '建筑、材料生产、装备加工、药剂、带出提升。'],
    ['配置表结构说明', 'Excel 配置表与程序读取边界。'],
])
mem.add_heading('5. 当前可用 skill 记录', level=1)
table(mem, ['skill', '适用场景'], [
    ['update-config', '修改 Claude Code 设置、权限、hooks、环境变量。'],
    ['keybindings-help', '调整 Claude Code 快捷键。'],
    ['simplify', '审查已改代码并进行复用/质量/效率优化。'],
    ['fewer-permission-prompts', '减少常见权限弹窗。'],
    ['loop', '按固定间隔执行任务。'],
    ['claude-api', 'Claude API / Anthropic SDK 相关开发。'],
    ['init', '初始化 CLAUDE.md。'],
    ['review', '审查 Pull Request。'],
    ['security-review', '进行安全审查。'],
])
mem.add_heading('6. 后续阅读建议', level=1)
num(mem, [
    '先阅读 SDC项目说明书.docx，快速理解产品边界。',
    '再阅读 SDC完整GDD初稿.docx，理解完整系统结构。',
    '进入 SDC项目开发Plan.docx，明确开发阶段与优先级。',
    '若要验证玩法，优先阅读 SDC第一版Demo说明书.docx。',
])
mem.save(base / 'SDC项目记忆与工作上下文.docx')

# 2. 产出文档总览
outdoc = Document()
setup(outdoc)
cover(outdoc, 'SDC项目产出文档总览', '用于汇总项目说明书、GDD、开发计划与 Demo 说明书等产出')
table(outdoc, ['文档项', '内容'], [
    ['文档定位', '产出文档索引'],
    ['主要用途', '统一存放并说明当前阶段所有正式产出文档'],
    ['更新方式', '每次新增正式产出后补充到本索引'],
])
outdoc.add_heading('1. 当前正式产出文档', level=1)
table(outdoc, ['文档名称', '定位', '状态', '说明'], [
    ['SDC项目说明书.docx', '项目整体框架说明', '已产出', '用于定义项目定位、核心循环、系统边界。'],
    ['SDC完整GDD初稿.docx', '完整 GDD 初稿', '已产出', '用于完整描述系统设计总纲。'],
    ['SDC项目开发Plan.docx', '开发阶段计划', '已产出', '基于 GDD 生成的模块拆分、优先级与开发阶段规划。'],
    ['SDC第一版Demo说明书.docx', '第一版 Demo 说明书', '已产出', '以关卡内循环验证为主，不接入外围养成。'],
    ['SDC完整GDD问题与框架确认稿.docx', '过程确认文档', '已产出', '保留两轮问题与结构确认过程，不作为最终正式说明。'],
])
outdoc.add_heading('2. 当前文档关系', level=1)
num(outdoc, [
    'SDC项目说明书.docx：定义项目的高层框架和边界。',
    'SDC完整GDD初稿.docx：在项目说明书基础上扩展为完整设计总纲。',
    'SDC项目开发Plan.docx：根据 GDD 生成开发阶段与模块优先级。',
    'SDC第一版Demo说明书.docx：从开发计划中抽取第一版 Demo 的验证目标与范围。',
])
outdoc.add_heading('3. 后续计划中的正式产出', level=1)
bullet(outdoc, [
    '战斗规则说明',
    '技能说明',
    '装备系统策划案',
    '关卡策划案',
    '玩家角色策划案',
    '怪物策划案',
    '家园系统策划案',
    '配置表结构说明',
])
outdoc.save(base / 'SDC项目产出文档总览.docx')

# 3. 项目开发计划
plan = Document()
setup(plan)
cover(plan, 'SDC项目开发Plan', '基于完整 GDD 的阶段拆分、模块优先级与开发建议')
table(plan, ['文档项', '内容'], [
    ['文档定位', '项目开发计划'],
    ['生成依据', 'SDC完整GDD初稿.docx'],
    ['目标', '明确模块优先级、阶段里程碑与第一版 Demo 的实现顺序'],
])
plan.add_heading('1. 开发目标理解', level=1)
p(plan, '当前项目应优先验证关卡内刷图、掉落、Build 成型和道中结算循环。在系统复杂度较高的情况下，开发计划需要遵循“关卡逻辑优先于外围养成”的原则。')
bullet(plan, [
    '优先验证：关卡推进、战斗、掉落、装备驱动技能变化、道中结算、Boss 奖励。',
    '延期验证：复杂家园、宠物深度、时装完整系统、长期经济、完整外围成长。',
])
plan.add_heading('2. 模块优先级', level=1)
table(plan, ['优先级', '模块', '说明'], [
    ['P0', '关卡逻辑', '章节进入、房间生成、三选一传送门、Boss 房、胜败结算。'],
    ['P0', '基础战斗框架', '移动、普攻、3 小技能、1 觉醒、闪避、索敌、受击与死亡。'],
    ['P0', '基础装备结构', '装备掉落、装备穿戴、品质、基础词条读取。'],
    ['P0', '装备裂变增益', '装备词条对技能和普攻的基础影响。'],
    ['P0', '道中结算系统', '翻牌奖励、道中商店、对外传输、全局胜败带出。'],
    ['P1', '自选难度 Tag 系统', '进入章节前选 Tag，影响奖励、事件权重和掉落规则。'],
    ['P1', '事件系统基础版', '普通战斗、精英、Boss、商人、休息、深渊。'],
    ['P1', '局内背包与拾取', '手动拾取、背包容量、筛选排序锁定、自动拾取扩展接口。'],
    ['P2', '基础家园结构', '带出数量提升、药剂制造、基础装备加工接口。'],
    ['P2', '职业 / 转职扩展', '多个职业路线、觉醒拓展。'],
    ['P3', '宠物、时装、复杂外围', '用于长期版本扩展。'],
])
plan.add_heading('3. 建议开发阶段', level=1)
plan.add_heading('3.1 阶段一：可玩原型', level=2)
bullet(plan, [
    '完成单章节单层或简化三层推进。',
    '完成玩家基础战斗与怪物基础行为。',
    '完成基础掉落、基础装备读取与穿戴。',
    '完成基础 Boss 战与胜败结算。',
])
plan.add_heading('3.2 阶段二：核心循环验证', level=2)
bullet(plan, [
    '接入三选一房间推进。',
    '接入翻牌奖励、道中商店、对外传输。',
    '接入 3 层结构与第 1 / 2 / 3 层的节奏定位。',
    '接入基础自选难度 Tag。',
    '验证玩家是否能感受到刷图、掉装、Build 成型和收益选择。',
])
plan.add_heading('3.3 阶段三：成长与拓展', level=2)
bullet(plan, [
    '扩展更多职业、装备池、事件池、深渊、隐藏 Boss。',
    '接入基础家园结构。',
    '补齐长期资源流转与更多外围系统。',
])
plan.add_heading('4. 当前最推荐的实施顺序', level=1)
num(plan, [
    '关卡推进框架。',
    '战斗控制与怪物基础。',
    '掉落、背包、装备读取。',
    '装备对技能/普攻的基础裂变。',
    'Boss 与道中结算。',
    '自选难度 Tag。',
    '事件扩展。',
    '基础家园接口。',
])
plan.add_heading('5. 与第一版 Demo 的关系', level=1)
p(plan, '第一版 Demo 只验证关卡内循环，不接入外围养成。因此项目开发 Plan 中的 P0 和部分 P1 模块，应优先服务于 Demo 可玩性，而不是完整项目内容量。')
plan.save(base / 'SDC项目开发Plan.docx')

# 4. 第一版 Demo 说明书
demo = Document()
setup(demo)
cover(demo, 'SDC第一版Demo说明书', '第一版 Demo 以关卡内循环验证为主，不接入外围养成')
table(demo, ['文档项', '内容'], [
    ['文档定位', '第一版 Demo 说明书'],
    ['核心目标', '验证关卡内循环'],
    ['不验证内容', '复杂家园、宠物、时装、长期外围养成'],
])
demo.add_heading('1. Demo 目标', level=1)
p(demo, '第一版 Demo 的目的不是验证完整项目的全部系统，而是验证“进入关卡 → 房间推进 → 战斗 → 掉落 → Build 变化 → Boss → 道中结算 / 传输 / 胜败带出”的核心关卡循环是否成立。')
bullet(demo, [
    '验证关卡推进是否顺畅。',
    '验证战斗是否具备基础爽感。',
    '验证装备掉落与技能裂变是否能形成 Build 感知。',
    '验证 Boss 奖励、翻牌、传输与结算是否能形成目标反馈。',
    '验证自选难度 Tag 是否能带来清晰的高风险高收益差异。',
])
demo.add_heading('2. Demo 范围', level=1)
table(demo, ['模块', 'Demo 范围'], [
    ['章节', '1 个章节。'],
    ['层数', '默认 3 层。'],
    ['房间推进', '哈迪斯-like 三选一推进，不可返回。'],
    ['房间类型', '普通战斗、精英战斗、Boss、商人、休息、深渊（可按实现情况裁剪）。'],
    ['角色', '优先 1 个基础职业路线，建议枪手。'],
    ['技能', '普攻 + 3 小技能 + 1 觉醒。'],
    ['怪物', '普通怪、精英怪、每层关底 Boss。'],
    ['装备', '基础品质、基础掉落、基础词条、基础技能裂变。'],
    ['结算', '翻牌奖励、道中商店、对外传输、全局胜败带出。'],
    ['难度系统', '接入基础自选难度 Tag。'],
])
demo.add_heading('3. Demo 不接入内容', level=1)
bullet(demo, [
    '复杂家园养成。',
    '宠物正式系统。',
    '时装正式系统。',
    '长期生产、真实时间资源循环。',
    '复杂职业矩阵与多章节内容量。',
])
demo.add_heading('4. Demo 重点验证玩法', level=1)
num(demo, [
    '玩家进入章节前选择少量难度 Tag。',
    '玩家进入房间，完成战斗或事件目标。',
    '玩家手动拾取装备和道具。',
    '装备改变普攻或技能表现，玩家感知 Build 差异。',
    '玩家在三选一传送门中做路线决策。',
    '玩家击败每层 Boss 后体验翻牌奖励与道中商店。',
    '玩家执行道具对外传输并理解其价值。',
    '玩家在最终胜利、失败、主动退出三种情况下看到不同结算结果。',
])
demo.add_heading('5. Demo 成功标准', level=1)
bullet(demo, [
    '玩家能完整完成一轮关卡内循环。',
    '玩家能明确感知到装备掉落和 Build 变化的价值。',
    '玩家能理解道中结算、传输和胜败带出的差异。',
    '玩家能理解难度 Tag 带来的风险和收益提升。',
    '整体体验证明该核心循环值得继续扩展外围系统。',
])
demo.add_heading('6. Demo 推荐实现顺序', level=1)
num(demo, [
    '房间推进与基础地图。',
    '玩家战斗与怪物基础行为。',
    '掉落、拾取、背包。',
    '基础装备与技能裂变。',
    'Boss 与道中结算。',
    '自选难度 Tag。',
    '商人 / 休息 / 深渊等扩展事件。',
])
demo.save(base / 'SDC第一版Demo说明书.docx')

print('done')
