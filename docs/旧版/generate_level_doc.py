from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

out = Path(r'c:/Users/loofnn/Desktop/sdc/策划案文档/关卡地图说明书/关卡说明书.docx')


def set_font(run, size=10.5, bold=False, color=None):
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def setup(doc):
    styles = doc.styles
    styles['Normal'].font.name = '宋体'
    styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    styles['Normal'].font.size = Pt(10.5)
    for name, size, color in [('Heading 1', 16, '1F4E79'), ('Heading 2', 13, '1F4E79'), ('Heading 3', 11, '000000')]:
        st = styles[name]
        st.font.name = '宋体'
        st._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
    for sec in doc.sections:
        sec.top_margin = Cm(2.2)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.4)
        sec.right_margin = Cm(2.4)


def p(doc, text='', style=None, size=10.5, bold=False, color=None, align=None):
    para = doc.add_paragraph(style=style)
    para.paragraph_format.line_spacing = 1.25
    if align:
        para.alignment = align
    r = para.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    return para


def bullet(doc, items):
    for item in items:
        para = doc.add_paragraph(style='List Bullet')
        para.paragraph_format.line_spacing = 1.18
        r = para.add_run(item)
        set_font(r)


def num(doc, items):
    for item in items:
        para = doc.add_paragraph(style='List Number')
        para.paragraph_format.line_spacing = 1.18
        r = para.add_run(item)
        set_font(r)


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run(h)
        set_font(r, bold=True)
        shade_cell(cell, 'D9EAF7')
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            para = cells[i].paragraphs[0]
            r = para.add_run(str(val))
            set_font(r, size=10)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return t


doc = Document()
setup(doc)

p(doc, '关卡说明书', size=22, bold=True, color='1F4E79', align=WD_ALIGN_PARAGRAPH.CENTER)
p(doc, 'SDC 关卡、房间与路线选择系统说明', size=12, color='595959', align=WD_ALIGN_PARAGRAPH.CENTER)
p(doc, '')
table(doc, ['文档项', '内容'], [
    ['文档名称', '关卡说明书'],
    ['所属系统', '关卡、房间与路线选择系统'],
    ['文档定位', '说明章节推进、房间结构、事件选择、Boss 流程与关卡内节奏'],
    ['写作边界', '方向清楚 + 关键规则，不展开具体表格字段与数值细节'],
])

doc.add_heading('1. 系统定位', level=1)
p(doc, '关卡系统是 SDC 的核心体验承载层，负责把战斗、掉落、Build 成长、事件选择、Boss 奖励和道中结算串联成完整的关卡内循环。')
bullet(doc, [
    '玩家通过章节进入关卡。',
    '玩家在房间推进中完成战斗与事件。',
    '玩家通过三选一传送门做路线选择。',
    '玩家在每层关底 Boss 后进入翻牌、商店、传输与继续推进的节奏。',
    '玩家最终在胜利、失败或主动退出中完成整轮结算。',
])

doc.add_heading('2. 关卡系统目标', level=1)
bullet(doc, [
    '提供清晰、稳定、可重复游玩的关卡内循环。',
    '为装备掉落、Build 成型与难度挑战提供结构化舞台。',
    '让玩家在每次房间选择中感受到风险与收益取舍。',
    '让每层 Boss、翻牌奖励、道中商店和传输机会形成阶段性高潮。',
    '支撑后续章节扩展、事件扩展、隐藏关卡扩展和高难挑战扩展。',
])

doc.add_heading('3. 章节结构', level=1)
table(doc, ['模块', '说明'], [
    ['章节数量目标', '当前整体目标为 3～5 个章节。'],
    ['章节差异', '每个章节拥有独立美术风格，并可配置不同怪物池、Boss 池、事件池、装备池。'],
    ['章节入口', '玩家从局外进入章节，并在章节页面选择自选难度 Tag。'],
    ['章节边界', '一个章节构成一轮完整探索，玩家在通关、失败或退出后结束本轮关卡。'],
])


doc.add_heading('4. 层结构', level=1)
p(doc, '一轮章节探索默认由 3 层组成，每层都围绕“房间推进 → 关底 Boss → 翻牌 / 商店 / 传输 → 下一层”展开。')
table(doc, ['层级', '定位', '说明'], [
    ['第 1 层', '成型层', '用于让玩家建立基础 Build、获取首轮关键掉落、熟悉当前局节奏。'],
    ['第 2 层', '强化层', '用于加深 Build 差异，强化玩家对路线、事件和收益的判断。'],
    ['第 3 层', '检验层', '用于检验当前 Build 与玩家操作，提供最终通关挑战。'],
])
p(doc, '每层关底都存在 1 个 Boss，Boss 击败后触发统一的层级奖励模板。')


doc.add_heading('5. 房间结构', level=1)
p(doc, '关卡以房间为基本推进单位。玩家进入房间后，需要完成战斗、事件或特定目标，房间完成后才能解锁下一步路线。')
table(doc, ['房间规则', '说明'], [
    ['单层房间数', '单层目标约 8～10 个房间，示例可在 9～11 个区间内调整。'],
    ['房间完成条件', '完成怪物清理、事件交互或对应目标后视为房间完成。'],
    ['房间奖励', '房间内主要通过怪物掉落、事件奖励、局内道具等方式给予收益。'],
    ['房间推进', '房间完成后出现下一步传送门选项。'],
])

doc.add_heading('6. 路线推进方式', level=1)
p(doc, '地图推进采用哈迪斯-like 结构：玩家完成当前房间后，进入下一次房间选择。')
bullet(doc, [
    '房间完成后通常出现 3 个可选传送门。',
    '玩家从三选一中决定下一房间。',
    '路线不可返回，该规则为绝对成立。',
    '地图不再使用迷雾预览机制。',
])
table(doc, ['玩家可见信息', '说明'], [
    ['房间类型', '例如普通战斗、精英、商人、休息、深渊等。'],
    ['奖励类型', '例如装备、局内道具、资源倾向等。'],
    ['危险等级', '帮助玩家预估挑战强度。'],
    ['事件图标', '用于快速识别功能型房间。'],
    ['特殊标记', '是否为精英、深渊或特殊事件。'],
])


doc.add_heading('7. 房间类型框架', level=1)
table(doc, ['房间类型', '定位', '一阶段建议'], [
    ['普通战斗', '基础清图、基础掉落、Build 成型的主要场景', '必须完成'],
    ['精英战斗', '更高挑战与更高奖励的过渡房间', '必须完成'],
    ['Boss 房', '每层关底挑战，承担节奏高潮与结算入口', '必须完成'],
    ['商人房', '出售、购买、局内资源调度', '必须完成'],
    ['休息房', '缓冲节奏、功能恢复或准备性事件', '建议完成'],
    ['深渊房', '高风险高收益、强难度验证', '建议完成'],
    ['NPC / 特殊剧情房', '角色、时装或特殊功能扩展入口', '后续扩展'],
])


doc.add_heading('8. 特殊传送门机制', level=1)
p(doc, '除常规三选一传送门外，关卡中还可额外触发特殊传送门。')
bullet(doc, [
    '特殊传送门与三选一传送门规则独立。',
    '特殊传送门指向特殊关卡库。',
    '特殊传送门的触发可受随机、事件权重、难度 Tag、章节进度、特定道具等多维度影响。',
    '若特殊关卡库只配置一个目标，则指向唯一关卡；若配置多个目标，则按权重随机。',
])
p(doc, '特殊传送门可以用于深渊、隐藏事件、特殊奖励、主题房间或指定 NPC 事件等扩展内容。')


doc.add_heading('9. 自选难度 Tag 与关卡关系', level=1)
p(doc, '玩家在进入章节前选择自选难度 Tag。Tag 不直接改变关卡结构本身，但会显著影响房间体验和关卡收益。')
bullet(doc, [
    '自选难度 Tag 会影响怪物数值、怪物机制、玩家限制等。',
    'Tag 会影响事件权重，例如深渊、特殊传送门或稀有事件出现概率。',
    'Tag 会影响掉落收益、Boss 保底、翻牌品质、商店内容等。',
    'Tag 形成关卡系数分，构成玩家的后期挑战目标。',
])


doc.add_heading('10. 每层 Boss 流程', level=1)
p(doc, '每层都以关底 Boss 作为阶段终点。Boss 被击败后，关卡进入统一的层级奖励流程。')
num(doc, [
    '击败当前层 Boss。',
    '触发翻牌奖励。',
    '进入道中商店。',
    '获得一次对外传输道具机会。',
    '玩家继续进入下一层，或在可退出节点按失败结算离开。',
])
p(doc, '最后一层 Boss 击败后，视为关卡胜利；同时存在小概率触发隐藏 Boss 的扩展规则。')


doc.add_heading('11. 翻牌、商店与传输', level=1)
table(doc, ['模块', '说明'], [
    ['翻牌奖励', '玩家从多张牌中选择 1 张，主要获得可带出局外的道具和装备。'],
    ['奖励品质', '铜牌、银牌、金牌影响奖励品质并读取不同奖励库。'],
    ['道中商店', '每层 Boss 后固定出现，用于出售、购买装备、材料和局内道具。'],
    ['对外传输', '玩家可将一定数量道具立即安全送入局外背包，不再受后续失败影响。'],
])
p(doc, '这套流程让每层 Boss 后都形成一次明确的阶段高潮，也是关卡内循环区别于单纯刷房的重要节奏点。')


doc.add_heading('12. 胜利、失败与退出', level=1)
bullet(doc, [
    '击败最后一层 Boss 视为关卡胜利。',
    '最后一层 Boss 后可能触发隐藏 Boss，玩家可选择挑战或直接按胜利结算。',
    '任意战斗失败均按失败结算。',
    '玩家可在道中主动退出，按失败结算。',
    '玩家可在层级奖励 / 商店阶段主动放弃，按失败带出。',
])
p(doc, '胜利和失败带出的数量为全局数值，不再随层级递增；其最终值由家园、天赋、章节奖励等系统共同影响。')


doc.add_heading('13. 关卡节奏建议', level=1)
table(doc, ['内容', '建议方向'], [
    ['普通房节奏', '前期约 30～90 秒，后期 Build 成型后可缩短至 10～30 秒。'],
    ['精英房节奏', '建议为普通房的 1.5～2 倍。'],
    ['Boss 房节奏', '建议约 2～4 分钟，最终以体验调试为准。'],
    ['整体感受', '节奏上强调边走边打、高频移动、掉落驱动继续推进。'],
])
p(doc, '关卡节奏设计的重点不是平均时长本身，而是确保玩家在房间推进中持续感到掉落、成长和选择的动力。')


doc.add_heading('14. 一阶段 Demo 对关卡系统的要求', level=1)
bullet(doc, [
    '至少完成 1 个章节、3 层结构。',
    '至少完成普通战斗、精英、Boss、商人四类核心房间。',
    '完成三选一房间推进与不可返回规则。',
    '完成每层 Boss 后的翻牌、商店、传输流程。',
    '完成全局胜利 / 失败 / 主动退出结算。',
    '至少接入一批基础自选难度 Tag。',
])


doc.add_heading('15. 与其他系统的接口关系', level=1)
table(doc, ['关联系统', '接口关系'], [
    ['战斗系统', '关卡房间决定怪物出现、战斗场景和挑战密度。'],
    ['装备系统', '关卡是装备掉落和 Build 成型的主要来源。'],
    ['掉落系统', '不同房间、Boss、深渊、Tag 权重共同决定掉落产出。'],
    ['事件系统', '房间类型和特殊传送门承载事件分布。'],
    ['自选难度 Tag 系统', '影响关卡收益、事件权重、掉落与挑战强度。'],
    ['道中结算系统', '每层 Boss 后进入翻牌、商店、传输与结算流程。'],
])


doc.add_heading('16. 后续专项文档衔接', level=1)
p(doc, '本说明书只定义关卡地图与推进系统的整体方向。具体配置和细节需要在后续专项文档中展开。')
table(doc, ['后续文档', '说明'], [
    ['关卡表结构说明', '章节、层、房间、事件权重、特殊传送门等配置规则。'],
    ['事件系统策划案', '商人、休息、深渊、NPC 等具体事件设计。'],
    ['怪物策划案', '不同房间和章节中的怪物内容。'],
    ['掉落系统策划案', '房间、Boss、深渊与 Tag 的掉落关系。'],
    ['战斗规则说明', '关卡推进中战斗节奏与角色表现的具体规则。'],
])

doc.save(out)
print(out.as_posix())
