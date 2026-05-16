from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

out = Path(r'c:/Users/loofnn/Desktop/sdc/策划案文档/关卡地图说明书/关卡说明书对齐问题.docx')


def set_font(run, size=10.5, bold=False, color=None):
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def setup(doc):
    styles = doc.styles
    styles['Normal'].font.name = '宋体'
    styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    styles['Normal'].font.size = Pt(10.5)
    for name, size, color in [('Heading 1', 16, '1F4E79'), ('Heading 2', 13, '1F4E79'), ('Heading 3', 11, '000000')]:
        st = styles[name]
        st.font.name = '宋体'
        st._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
    for sec in doc.sections:
        sec.top_margin = Cm(2.2)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.4)
        sec.right_margin = Cm(2.4)


def p(doc, text='', style=None, size=10.5, bold=False, color=None, align=None):
    para = doc.add_paragraph(style=style)
    para.paragraph_format.line_spacing = 1.25
    if align:
        para.alignment = align
    r = para.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    return para


def bullet(doc, items):
    for item in items:
        para = doc.add_paragraph(style='List Bullet')
        para.paragraph_format.line_spacing = 1.18
        r = para.add_run(item)
        set_font(r)


def num(doc, items):
    for item in items:
        para = doc.add_paragraph(style='List Number')
        para.paragraph_format.line_spacing = 1.18
        r = para.add_run(item)
        set_font(r)


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run(h)
        set_font(r, bold=True)
        shade_cell(cell, 'D9EAF7')
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            para = cells[i].paragraphs[0]
            r = para.add_run(str(val))
            set_font(r, size=10)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return t


def reply_box(doc, label='你的回复 / 修改意见：'):
    p(doc, label, bold=True, color='C00000')
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    cell = t.rows[0].cells[0]
    cell.text = '\n\n\n'
    shade_cell(cell, 'FFF2CC')
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def q(doc, text):
    p(doc, text, bold=True)
    reply_box(doc)


doc = Document()
setup(doc)

p(doc, '关卡说明书对齐问题', size=22, bold=True, color='1F4E79', align=WD_ALIGN_PARAGRAPH.CENTER)
p(doc, '基于参考策划案的理解、问题、阻力与待确认项', size=12, color='595959', align=WD_ALIGN_PARAGRAPH.CENTER)
p(doc, '')
table(doc, ['文档项', '内容'], [
    ['文档定位', '对齐用文档'],
    ['用途', '说明我对参考策划案的理解、我认为参考案的优点/问题、以及以它为模板推进 SDC 关卡说明书时的阻力与待确认问题'],
    ['沟通方式', '请直接在黄色区域中填写回复、修改意见或补充说明'],
])

doc.add_heading('一、我对“真实项目使用标准策划案”的理解', level=1)
p(doc, '我阅读了你提供的两个参考策划案后，当前理解是：你要的并不是“产品层面的高层说明文档”，而是“可以直接进入真实项目协作、可供程序/UI/数值/测试落地执行的专项策划案”。')
bullet(doc, [
    '它必须同时服务多个岗位：策划、程序、UI、美术、测试。',
    '它不能只讲方向，还必须讲规则、流程、交互、表现、配置表、边界条件。',
    '它不仅要回答“系统是什么”，还要回答“玩家怎么点、客户端怎么表现、程序怎么判定、表格怎么支持、美术要做什么”。',
    '它必须带有可配置的数据结构意识，而不是纯文本描述。',
])
p(doc, '基于两个参考案，我理解一个达到真实项目使用标准的专项策划案，至少需要包含这几层内容：')
num(doc, [
    '文档总览：责任人、状态、修订记录、设计目的、跟进与验收流程。',
    '功能规则：系统定位、设计目的、玩法概述、世界观包装、规则说明。',
    '交互逻辑：页面结构、状态机、点击行为、异常提示、跳转规则。',
    '客户端表现：界面展示、按钮状态、图标、弹窗、飘字、特殊表现。',
    '程序规则：判定条件、流程节点、异常分支、优先级、互斥关系。',
    '表格结构：输出文件名、字段、字段类型、配置样例、读表关系。',
    '协作接口：哪些内容给 UI、美术、程序、测试使用。',
])


doc.add_heading('二、我如何理解这两个参考策划案', level=1)
doc.add_heading('1. 对《绝境突围功能文档》的理解', level=2)
bullet(doc, [
    '这是一个“玩法系统型”策划案。',
    '它不是只描述玩法概念，而是把玩法入口、开放条件、核心流程、层开放规则、事件类型、奖励逻辑、格子事件表、事件表、增益表都拉进同一个文档框架。',
    '它强调“玩法规则 + 配置表结构”并行书写，让程序和策划都知道规则最终如何落在表里。',
    '它明显面向真实协作场景：程序能知道要读哪些表，策划能知道要配什么，美术/UI 也能知道有哪些界面和事件需要支撑。',
])
doc.add_heading('2. 对《装备系统文档 v1.1》的理解', level=2)
bullet(doc, [
    '这是一个“系统养成型”策划案。',
    '它的结构比高层 GDD 更细，已经进入“系统功能可拆实现”的层级。',
    '它把“功能规则&表格”“交互逻辑”“装备表”“装备强化表”拆页处理，说明专项策划案不是一段长文，而是“规则 + 交互 + 数据结构”的组合。',
    '它很强调交互状态，例如空位态、装备态、未解锁态、筛选展开态，这意味着真实项目策划案必须把状态机写出来。',
    '它的表格不是附录，而是功能本体的一部分。装备字段、强化消耗、改造限制，本身就是规则的一部分。',
])

doc.add_heading('三、参考策划案对我最大的启发', level=1)
bullet(doc, [
    '专项策划案不能只讲“系统有什么”，必须讲“每个环节具体怎么跑”。',
    '文档必须考虑程序的漏判风险，所以要补齐边界条件、异常状态、状态切换。',
    '文档必须考虑客户端表现，不然 UI 和前端无法知道页面、弹窗、按钮、提示如何呈现。',
    '文档必须和配置表联动，不然策划的规则和程序的实现之间会断层。',
    '专项策划案需要天然带着“这是给团队执行的”意识，而不是“这是给人阅读理解的”意识。',
])


doc.add_heading('四、我认为参考策划案本身存在的问题', level=1)
p(doc, '参考策划案整体方向是对的，也明显比我当前交付的《关卡说明书》更接近真实项目标准。但从“作为模板直接复用”的角度，我也看到了几个问题。')
bullet(doc, [
    '问题 1：部分规则仍然偏散，规则文本、交互逻辑、表格结构之间的映射关系没有完全显式写清。',
    '问题 2：虽然有交互逻辑页签，但“客户端表现”和“程序判定”仍然可能交叉，需要更明确区分“前端表现规则”和“服务端/逻辑判定规则”。',
    '问题 3：表格结构里有样例，但部分字段解释不足。如果新人接手，可能仍需要口头补充。',
    '问题 4：美术需求在参考案里仍不够独立。它更偏 UI/交互，但纯美术资源需求、场景需求、特效需求没有完全系统化列出来。',
    '问题 5：流程图、状态图在当前导出的参考内容里不够直观。如果没有配图或脑图页签，复杂玩法的理解成本仍然高。',
])
p(doc, '换句话说：参考策划案已经是“真实项目可用”的方向，但如果要拿它做 SDC 的统一模板，我认为还可以再补强“状态机分层、程序边界、表格字段说明、美术资源清单”四块。')


doc.add_heading('五、以参考策划案为模板时，我完成 SDC 关卡说明书的阻力是什么', level=1)
p(doc, '当前最大的阻力，不是我不能写更多，而是我缺少足够多的“执行级别输入信息”。如果要按参考策划案标准写 SDC 关卡说明书，我需要的不再只是产品概念，而是更接近落地的规则输入。')
table(doc, ['阻力点', '具体表现'], [
    ['交互输入不足', '我知道有三选一传送门、翻牌、商店、传输，但不知道每一步页面怎么展示、按钮怎么点、是否有确认弹窗、错误提示是什么。'],
    ['程序规则输入不足', '我知道有胜利/失败/主动退出，但不知道各种边界情况的优先级和完整判定顺序。'],
    ['表格结构输入不足', '我知道需要表格，但目前没有“关卡系统最终要拆成哪些表、表之间怎么关联、字段怎么命名”的明确约定。'],
    ['美术需求输入不足', '我知道要有房间、传送门、章节、美术风格，但不知道要不要把场景资源、UI 资源、特效资源、icon 需求单独列项。'],
    ['系统边界不足', '我还不知道“关卡说明书”这份专项文档应该写到多细，哪些写在关卡说明书里，哪些拆给事件系统、掉落系统、怪物系统。'],
])


doc.add_heading('六、如果以参考案为模板，我会如何重构 SDC 的《关卡说明书》', level=1)
p(doc, '如果按你给的参考案标准重写，我理解新的《关卡说明书》至少应拆成以下结构：')
num(doc, [
    '文档总览：责任人、状态、修订记录、设计目的。',
    '功能规则：关卡系统定位、开放条件、章节结构、层结构、房间结构、三选一路线、Boss 流程、结算规则。',
    '流程图：章节进入流程、房间推进流程、Boss 后流程、胜败结算流程、特殊传送门流程。',
    '交互逻辑：章节界面、Tag 选择界面、房间选择界面、翻牌界面、商店界面、传输界面、失败结算界面。',
    '程序规则：事件生成、房间刷新、特殊传送门触发、路线锁定、隐藏 Boss 判定、退出判定。',
    '表格结构：章节表、层配置表、房间池表、事件权重表、特殊传送门表、Boss 奖励表、翻牌奖励表、传输表、Tag 表。',
    '客户端表现与 UI 需求：页面、icon、按钮、弹窗、状态提示、飘字、奖励表现。',
    '美术需求：章节场景、房间底图、传送门、Boss 房表现、深渊表现、特殊房表现、图标与框体需求。',
])


doc.add_heading('七、我当前不理解或需要你确认的点', level=1)
q(doc, '1. 你希望 SDC 的专项策划案格式，更接近《绝境突围功能文档》这种“玩法系统 + 配表结构”的模板，还是更接近《装备系统文档》这种“规则 + 交互 + 表格拆页”的模板？还是两者结合？')
q(doc, '2. 对于《关卡说明书》，你希望它是否独立承担这些内容：规则逻辑、交互逻辑、客户端表现、表格结构说明、美术需求？还是其中一部分应该拆给别的文档？')
q(doc, '3. 关卡系统的“交互逻辑”你希望细到什么程度？例如：按钮文案、点击后弹窗、灰态/解锁态、错误提示、返回逻辑，是否都要写进专项策划案？')
q(doc, '4. 关卡系统的“程序规则”你希望细到什么程度？例如：进入房间时先判定什么、击败 Boss 后先结算什么、特殊传送门和三选一的优先级、主动退出和隐藏 Boss 的优先级，是否都要明确？')
q(doc, '5. 关卡系统的“联合表格结构说明”你希望我做到什么层级？是只列出需要哪些表和字段方向，还是要像参考案一样，直接写出表名、输出文件名、字段名、字段类型和配置样例？')
q(doc, '6. 关卡系统的“美术需求”你希望写到哪一层？仅列资源清单，还是要细到每个界面、每种房间、每类传送门、每种图标和特效需求？')
q(doc, '7. 关卡说明书和这些专项文档的边界，你希望怎么划分：事件系统策划案、掉落系统策划案、怪物策划案、战斗规则说明？例如“商店规则”是写在关卡说明书里，还是写在事件系统里，关卡说明书只保留流程入口？')
q(doc, '8. 参考策划案中，你最希望我优先学会并复用的部分是什么？是文档总览结构、交互逻辑写法、表格结构写法，还是“规则 + 表格联动”的方法？')
q(doc, '9. 你希望我下一步做的是：先重写《关卡说明书》结构模板，再补规则；还是先通过新一轮问题把关卡系统补细，再重写成真实项目标准版本？')


doc.add_heading('八、我当前的结论', level=1)
p(doc, '我接受你对当前《关卡说明书》交付质量“不合格”的判断，而且我认同这个判断。问题不在于文档有没有内容，而在于它仍然是“高层设计说明”，不是“真实项目专项策划案”。')
p(doc, '如果接下来要达到你给的参考标准，我会切换写法：从“系统介绍式写法”切到“规则 / 交互 / 程序 / 表格 / 美术联合说明写法”。但在动笔之前，我需要你先帮我确认上面这些边界和粒度问题。')
p(doc, '我建议：你先在本 Word 中回复我上面的关键问题。确认后，我再以你认可的模板，重写 SDC 的《关卡说明书》。')

doc.save(out)
print(out.as_posix())
