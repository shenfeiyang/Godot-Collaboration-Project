from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

out = Path(r'c:/Users/loofnn/Desktop/sdc/SDC项目说明书.docx')

def set_font(run, size=10.5, bold=False, color=None):
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def set_doc_defaults(doc):
    styles = doc.styles
    styles['Normal'].font.name = '宋体'
    styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    styles['Normal'].font.size = Pt(10.5)
    for name, size, color in [('Heading 1', 16, '1F4E79'), ('Heading 2', 14, '1F4E79'), ('Heading 3', 12, '000000')]:
        style = styles[name]
        style.font.name = '宋体'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)

def add_para(doc, text='', style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_font(r)
    return p

def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.line_spacing = 1.2
        r = p.add_run(item)
        set_font(r)

def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.line_spacing = 1.2
        r = p.add_run(item)
        set_font(r)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, bold=True)
        set_cell_shading(cell, 'D9EAF7')
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_font(r, size=10)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return table

doc = Document()
set_doc_defaults(doc)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('SDC项目说明书')
set_font(r, size=24, bold=True, color='1F4E79')
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run('第一版｜产品整体框架说明与系统边界声明')
set_font(r, size=12, color='595959')

add_para(doc, '')
add_table(doc, ['文档项', '内容'], [
    ['文档名称', 'SDC项目说明书'],
    ['文档定位', '产品整体框架说明 + 系统边界声明'],
    ['适用阶段', '项目前期设计梳理、程序理解项目基础框架'],
    ['覆盖范围', '项目定位、核心体验、基础玩法循环、系统框架、后续文档拆分'],
    ['不覆盖范围', '第一版 Demo 开发细节、单系统详细机制、具体数值表、技能/装备完整配置表'],
])

add_para(doc, '')
doc.add_heading('修改记录', level=1)
add_table(doc, ['版本号', '更新时间', '更新内容', '修改者'], [
    ['V 1.0.0', '2026年5月12日', '第一版项目说明书，用于确认产品整体框架与系统边界。', ''],
])

add_para(doc, '')
doc.add_heading('目录', level=1)
add_bullets(doc, [
    '一：项目概述', '二：核心体验目标', '三：基础玩法循环', '四：关卡与层结构',
    '五：层结算与带出机制', '六：战斗框架概述', '七：装备系统框架概述', '八：技能与 Build 框架概述',
    '九：疲劳值机制概述', '十：事件系统框架概述', '十一：局外养成框架', '十二：美术与题材方向',
    '十三：后续文档拆分', '十四：整体文档定位',
])

doc.add_page_break()

doc.add_heading('一：项目概述', level=1)
add_para(doc, 'SDC 是一款像素风、2D 俯视角、暗黑-like 刷宝肉鸽 ARPG。')
add_para(doc, '项目核心围绕“刷装备、构筑 Build、推进多层关卡、在风险收益中选择撤离或继续挑战”展开。玩家通过长期刷取装备、组合词条与套装、解锁职业与觉醒分支，不断形成新的技能表现和战斗流派。')
add_para(doc, '战斗体验偏弹幕躲避与技能循环，装备长期驱动角色成长与技能裂变。整体参考方向包括 DNF 式刷图掉装、暗黑破坏神式装备流派构筑，以及带有肉鸽风险收益选择的多层探索结构。')
add_table(doc, ['项目维度', '说明'], [
    ['类型定位', '像素俯视角、暗黑-like 刷宝、肉鸽 ARPG'],
    ['核心驱动', '装备刷宝、Build 构筑、技能裂变、层推进与撤离选择'],
    ['主要视角', '2D 俯视角'],
    ['美术方向', '像素风、奇幻地下城、暗黑地牢氛围'],
    ['核心参考', 'DNF 刷图掉装体验、暗黑式装备流派、肉鸽式风险收益选择'],
])

doc.add_heading('二：核心体验目标', level=1)
add_para(doc, '本项目希望玩家持续游玩的主要原因，是通过一次次关卡探索获得装备、材料和稀有构筑资源，并将这些资源沉淀为长期角色成长。')
add_bullets(doc, [
    '刷装备带来的品质反馈：怪物和 Boss 掉落装备，稀有品质和关键词条形成明确追求目标。',
    'Build 构筑带来的成长反馈：装备词条、套装、SP 技能、转职、觉醒共同改变技能表现和战斗方式。',
    '层结算带来的风险收益：每层结束后，玩家需要在“带出收益”与“继续挑战更高奖励”之间选择。',
    '多层推进带来的挑战反馈：越深入关卡，难度越高，奖励越丰富，Boss 掉落和稀有装备追求更强。',
    '永久养成带来的积累反馈：失败只降低本次可带出的收益，不摧毁玩家已经获得的长期成长。',
])

doc.add_heading('三：基础玩法循环', level=1)
add_para(doc, '基础玩法循环是程序理解项目整体骨架的核心。本项目的一轮探索由局外准备、局内推进、层末结算、战利品带出和局外成长组成。')
add_numbered(doc, [
    '玩家从家园或局外准备界面进入关卡。',
    '进入章节关卡，按照房间节点进行推进。',
    '清理普通战斗、精英战斗、商店、休息区等不同事件类型房间。',
    '获得装备、材料、技能相关道具等战利品。',
    '每进入房间消耗疲劳值，路线选择受到疲劳、奖励和风险影响。',
    '到达当前层末 Boss 房。',
    '击败 Boss 后进入层结算。',
    '玩家选择带出部分战利品返回家园，或继续挑战下一层。继续挑战后，难度提高、奖励更好，但累计战利品继续承担失败风险。',
    '失败时按所在层失败带出数量结算。',
    '带回的装备和资源进入局外长期养成。',
])
add_para(doc, '该循环的重点不是一次性通关，而是在“刷取—取舍—带出—养成—再挑战”的过程中形成长期追求。')

doc.add_heading('四：关卡与层结构', level=1)
add_para(doc, '关卡底层可采用格子或节点配置生成，但玩家实际体验应接近 DNF 式房间推进：清理当前房间后，通过传送门前往下一个节点。')
add_bullets(doc, [
    '地图底层支持格子或节点配置，具体行列参数可根据体验调整。',
    '玩家在清理房间后，通过 1～3 个传送门选择不同路线。',
    '路线选择可受到事件类型、奖励类型、危险等级和疲劳值影响。',
    '存在迷雾预览机制，玩家可看到下一节点的事件类型、奖励类型、危险等级等信息。',
    '默认一轮探索为 3 层。未来可扩展无尽层、隐藏层、低概率第四层等长期挑战内容。',
    '每层末尾为 Boss 房。',
    '第一阶段房间类型框架可先定义为普通战斗、精英战斗、Boss；商人、休息、深渊等事件作为后续扩展系统保留接口。',
])

doc.add_heading('五：层结算与带出机制', level=1)
add_para(doc, '层结算与带出机制是本项目最重要的风险收益系统之一。玩家每完成一层后，不是简单进入下一层，而是需要判断当前收益是否值得保留，或是否继续挑战更高难度。')
add_bullets(doc, [
    '每层结束后出现结算选择。',
    '可选行为包括：带出、继续、出售。',
    '带出格子为装备与道具共用格子，玩家需要在装备、材料、技能相关道具之间取舍。',
    '继续挑战后，上一层未带出的战利品继续承担失败风险。',
    '失败不损失永久养成，只降低本轮可带出的收益数量。',
    '家园建筑可提升带出格子数量。',
])
add_table(doc, ['层数', '成功带出数量', '失败带出数量', '说明'], [
    ['第一层', '15', '8', '示例参数，最终以配置表为准'],
    ['第二层', '25', '15', '示例参数，最终以配置表为准'],
    ['第三层', '35', '25', '示例参数，最终以配置表为准'],
])
add_para(doc, '该机制用于避免玩家只在最终通关时才结算收益，使每层末尾都形成一次明确的风险判断。')

doc.add_heading('六：战斗框架概述', level=1)
add_para(doc, '战斗为 2D 俯视角弹幕压力战斗。玩家需要通过移动、普攻、小技能、大招、闪避和切目标完成战斗。')
add_table(doc, ['操作/能力', '定位'], [
    ['移动', '躲避弹幕、调整站位、寻找输出窗口。'],
    ['普攻', '由武器决定攻击方式，几乎无冷却、低消耗、高频低伤害。'],
    ['小技能', '承担主要战斗循环，有一定冷却和资源消耗，效果强于普攻。'],
    ['大招', '长冷却、高强度，用于关键爆发和 Boss 战。'],
    ['闪避', '具有无敌帧、小位移、体力消耗，可被装备词条强化。'],
    ['切目标', '自动索敌为主，切目标作为辅助功能，解决玩家特定索敌需求。'],
])
add_para(doc, 'Boss 战承担操作检验、Build 检验和掉落高潮。具体怪物、Boss 阶段、技能参数、弹幕规则等内容不在本项目说明书展开，后续进入《战斗说明文档》和《怪物与 Boss 设计文档》。')

doc.add_heading('七：装备系统框架概述', level=1)
add_para(doc, '装备是本项目第一核心系统。装备承担数值成长、流派构筑、技能裂变和长期刷宝目标。')
add_table(doc, ['装备分区', '部位', '定位'], [
    ['武器', '武器', '决定攻击方式，拥有特殊技能库，输出能力强。'],
    ['服饰', '头盔、上衣、腰带、裤子、鞋子', '偏生存和基础属性，可形成服饰套装。'],
    ['饰品', '戒指、项链、手环', '偏均衡和特殊构筑，可形成饰品套装。'],
])
add_para(doc, '装备品质分为白、绿、蓝、紫、粉、橙、红。不同品质承担不同的成长阶段和追求价值。低品质装备也需要具备价值，可通过出售、融合、提取稀有词条等方式进入循环。')
add_bullets(doc, [
    '白、绿：低价值装备，可出售局内货币，也可作为融合材料。',
    '蓝、紫：过渡装备，同时可能产出高价值词条，可用于词条提取或融合。',
    '粉：中前期强力装备，部分可在后期保留。',
    '橙：毕业装备，作为长期追求目标。',
    '红：特殊毕业装备，获取难度高，周期长，作为稀有终局追求。',
])
add_para(doc, '装备包含品质、基础属性、二级词条、技能词条、套装、SP 技能等结构。词条提取/镶嵌、融合、洗练、分解、出售等系统在本说明书中只做框架占位，详细规则进入《装备系统策划案》。')

doc.add_heading('八：技能与 Build 框架概述', level=1)
add_para(doc, '技能系统不只由角色自身决定，而是由角色基础技能、装备、词条、套装、SP 技能、转职、觉醒、技能书等共同影响。装备是技能裂变的主要驱动力之一。')
add_bullets(doc, [
    '角色基础技能提供初始战斗框架。',
    '武器决定基础攻击方式，并提供特殊技能库。',
    '装备词条可改变技能数值、范围、冷却、弹道、分裂、穿透、追踪等表现。',
    '套装效果可围绕元素、技能类型、召唤物或特定流派进行强化。',
    'SP 技能用于对单个技能进行质变强化，是高品质装备的重要卖点。',
    '转职改变技能池、成长方向和流派方向，玩家在局外选择职业。',
    '觉醒永久解锁觉醒分支，并大幅强化技能体系。',
    '技能形态切换书作为后续系统，用于解锁或切换技能形态。',
])
add_para(doc, '本项目说明书只概括技能框架，不展开具体技能参数、技能样例和完整技能树。相关内容后续进入《技能系统策划案》和《战斗说明文档》。')

doc.add_heading('九：疲劳值机制概述', level=1)
add_para(doc, '疲劳值用于限制探索、制造风险收益、控制单轮时长，并鼓励玩家在路线中做出选择。')
add_bullets(doc, [
    '每进入一个房间消耗 1 点疲劳值。',
    '疲劳值阈值、增益和减益由配置表控制。',
    '高疲劳可获得正向增益，例如移动速度、攻击等方向的加成。',
    '低疲劳或 0 疲劳可获得减益，用于提高继续探索的压力。',
    '疲劳可通过局内事件、Boss 阶段奖励、疲劳药剂等方式恢复。',
    '疲劳药剂可在局内商店购买，后续也可由家园系统制造。',
])
add_para(doc, '疲劳值的具体数值、阈值和效果不在项目说明书展开，后续进入配置表和对应系统文档。')

doc.add_heading('十：事件系统框架概述', level=1)
add_para(doc, '事件系统用于丰富路线选择、资源补给、风险收益和长期拓展。本项目说明书只定义事件系统存在及扩展方向，不展开单个事件的详细规则。')
add_bullets(doc, [
    '事件可包含商人、休息、深渊、时装、角色、宠物等。',
    '商人可承担局内购买、出售和补给入口。',
    '休息事件可承担疲劳恢复、生命恢复或战斗准备功能。',
    '深渊事件可作为高风险高收益方向保留，具体形式后续确定。',
    '事件系统需要支持后续扩展和组合。',
    '第一阶段可先保留普通战斗、精英战斗、Boss 作为最小关卡框架。',
])
add_para(doc, '商人、休息、深渊等具体设计进入后续《事件系统策划案》。')

doc.add_heading('十一：局外养成框架', level=1)
add_para(doc, '局外养成用于承接玩家从关卡中带出的装备、材料和解锁资源，使玩家形成长期目标。')
add_bullets(doc, [
    '家园、角色、宠物、时装、章节推进等均作为长期拓展系统。',
    '当前优先级为：装备系统、时装系统、章节；其他内容后续版本扩展。',
    '家园系统至少承担部分局外成长功能，例如提升带出格子数、制造疲劳药剂等。',
    '角色系统、宠物系统、时装系统的具体玩法后续通过独立系统策划案说明。',
])
add_para(doc, '本项目说明书只说明局外养成的存在和方向，不展开具体功能细节。')

doc.add_heading('十二：美术与题材方向', level=1)
add_bullets(doc, [
    '美术方向：像素风。',
    '题材方向：奇幻地下城。',
    '氛围方向：偏暗黑地牢。',
    '具体角色、怪物、装备命名和美术表现由后续资源与美术方向决定。',
])
add_para(doc, '项目整体视觉应服务于刷宝、地下城探索、稀有装备反馈和弹幕战斗可读性。')

doc.add_heading('十三：后续文档拆分', level=1)
add_para(doc, '本项目说明书不承担完整 GDD 和所有系统细节的职责。为避免程序误解文档边界，后续需要将具体系统拆分为独立策划文档。')
add_table(doc, ['后续文档', '主要内容'], [
    ['《一阶段项目开发说明》', '第一版 Demo 的实际开发范围、功能拆分、基础实现目标。'],
    ['《装备系统策划案》', '装备品质、词条、套装、SP 技能、融合、洗练、提取、镶嵌等详细规则。'],
    ['《技能系统策划案》', '技能来源、技能树、技能裂变、技能书、转职与觉醒规则。'],
    ['《战斗说明文档》', '玩家操作、怪物行为、弹幕规则、伤害机制、战斗表现。'],
    ['《关卡地图系统策划案》', '节点生成、路线选择、迷雾预览、层结构、房间配置。'],
    ['《事件系统策划案》', '商人、休息、深渊、高风险奖励等事件规则。'],
    ['《家园系统策划案》', '家园建筑、带出格子提升、疲劳药剂制造等局外功能。'],
    ['《时装系统策划案》', '时装获取、展示、属性或套装方向。'],
    ['《怪物与 Boss 设计文档》', '怪物类型、Boss 机制、阶段设计、掉落规则。'],
    ['《配置表结构说明》', '各系统配置表字段、参数、引用关系和维护规则。'],
])

doc.add_heading('十四：整体文档定位', level=1)
add_para(doc, '《项目说明书》不是完整 GDD，也不是 Demo 开发说明，而是“产品整体框架说明 + 系统边界声明”。')
add_para(doc, '它主要用于让自己和程序理解：')
add_bullets(doc, [
    '这是一个什么类型的游戏。',
    '项目的核心循环是什么。',
    '哪些系统一定存在。',
    '哪些系统只是框架占位。',
    '哪些内容后续会拆成单独文档。',
    '早期程序架构需要为哪些系统预留接口。',
])
add_para(doc, '因此，本说明书中出现的装备、技能、事件、疲劳、家园等内容，主要用于确定系统方向和边界。具体数值、功能细节、配置表、单个技能或装备设计，应以后续单系统策划案为准。')

doc.save(out)
print(out.as_posix())
