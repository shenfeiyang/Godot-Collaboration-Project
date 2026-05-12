from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

out = Path(r'c:/Users/loofnn/Desktop/sdc/SDC完整GDD问题与框架确认稿.docx')


def set_font(run, size=10.5, bold=False, color=None):
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def set_doc_defaults(doc):
    styles = doc.styles
    styles['Normal'].font.name = '宋体'
    styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    styles['Normal'].font.size = Pt(10.5)
    for name, size, color in [('Heading 1', 16, '1F4E79'), ('Heading 2', 13, '1F4E79'), ('Heading 3', 11, '000000')]:
        style = styles[name]
        style.font.name = '宋体'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)


def p(doc, text='', style=None, size=10.5, bold=False, color=None, align=None):
    para = doc.add_paragraph(style=style)
    para.paragraph_format.line_spacing = 1.25
    if align:
        para.alignment = align
    run = para.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return para


def bullet(doc, items):
    for item in items:
        para = doc.add_paragraph(style='List Bullet')
        para.paragraph_format.line_spacing = 1.18
        run = para.add_run(item)
        set_font(run)


def num(doc, items):
    for item in items:
        para = doc.add_paragraph(style='List Number')
        para.paragraph_format.line_spacing = 1.18
        run = para.add_run(item)
        set_font(run)


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(h)
        set_font(run, bold=True)
        shade_cell(cell, 'D9EAF7')
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            para = cells[i].paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(str(val))
            set_font(run, size=10)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return t


def reply_box(doc, label='你的回复 / 修改意见：'):
    p(doc, label, bold=True, color='C00000')
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    cell = t.rows[0].cells[0]
    cell.text = '\n\n\n'
    shade_cell(cell, 'FFF2CC')
    return t


def section_questions(doc, title, questions):
    doc.add_heading(title, level=2)
    for q in questions:
        p(doc, q, bold=True)
        reply_box(doc)


doc = Document()
set_doc_defaults(doc)

p(doc, 'SDC 完整 GDD 问题与框架确认稿', size=22, bold=True, color='1F4E79', align=WD_ALIGN_PARAGRAPH.CENTER)
p(doc, '用途：用于对齐完整 GDD 的信息需求、章节结构、系统边界与后续推进顺序。请直接在黄色区域中填写回复、修改意见或“待定”。', size=11, color='595959', align=WD_ALIGN_PARAGRAPH.CENTER)

p(doc, '')
table(doc, ['文档项', '说明'], [
    ['文档名称', 'SDC 完整 GDD 问题与框架确认稿'],
    ['当前用途', '收集完整 GDD 所需信息，并确认 GDD 框架'],
    ['回复方式', '请直接在每个问题下方黄色区域填写回复或修改意见'],
    ['当前阶段', '对齐思路，不直接产出完整 GDD 正文'],
    ['后续动作', '根据你的回复，整理完整 GDD 目录与逐章写作顺序'],
])

p(doc, '')
doc.add_heading('一、需要优先确认的总问题', level=1)
section_questions(doc, '1. GDD 体量与用途', [
    '1.1 这个完整 GDD 最终主要给谁看？例如：自己梳理、程序开发、美术协作、团队成员、外部合作方。',
    '1.2 你希望完整 GDD 偏“开发落地型”，还是偏“设计总纲型”？或者两者结合？',
    '1.3 完整 GDD 是否需要覆盖到可直接拆任务的程度？还是只做到系统设计清楚即可？',
    '1.4 是否允许 GDD 中保留“待定项 / 后续策划案展开项”？哪些模块可以先待定？',
])
section_questions(doc, '2. 产品定位', [
    '2.1 游戏最终希望做成什么体量？小体量独立游戏、中体量买断制、长线更新、类手游长线养成，还是其他？',
    '2.2 首发平台是什么？PC / Steam / 手机 / Web / 主机？是否要同时考虑键鼠、手柄、触屏？',
    '2.3 商业模式倾向是什么？买断制、免费、DLC、资料片、外观收集，还是第一阶段完全不讨论？',
    '2.4 目标玩家是谁？DNF 刷宝玩家、暗黑/POE 构筑玩家、元气骑士前传玩家、哈迪斯肉鸽玩家、像素地下城玩家，还是混合？',
    '2.5 游戏整体偏硬核还是偏轻度？是否允许玩家研究复杂词条、抄 Build、刷毕业？',
])
section_questions(doc, '3. 设计支柱', [
    '3.1 当前设计支柱是否确认为：装备刷宝、装备驱动技能裂变、层结算撤离风险收益、弹幕压力+技能循环、永久养成不因失败摧毁？',
    '3.2 如果只能保留 3 个最核心支柱，你会保留哪 3 个？',
    '3.3 如果开发资源不足，哪些系统必须保留，哪些可以延期？',
])

p(doc, '')
doc.add_heading('二、完整 GDD 需要补齐的问题', level=1)
section_questions(doc, '4. 核心循环', [
    '4.1 秒级循环中，玩家最频繁做的事情是什么？移动、躲弹幕、普攻、技能、拾取、闪避、切目标，哪些最重要？',
    '4.2 战斗更偏“边走边打”，还是“找窗口爆发”？',
    '4.3 普攻是手动按键、按住连发，还是自动攻击？',
    '4.4 拾取是自动拾取、手动拾取，还是不同品质不同规则？',
    '4.5 单个普通房、精英房、Boss 房的目标时长大概是多少？',
    '4.6 一轮完整探索是否仍以 40 分钟～1.5 小时为目标？如果流程较长，是否需要中途保存？',
    '4.7 玩家主动退出、关闭游戏、战斗中退出时如何结算？',
])
section_questions(doc, '5. 战斗系统', [
    '5.1 最终技能槽位是“普攻 + 2 小技能 + 大招”，还是“普攻 + 3 小技能 + 大招”？',
    '5.2 是否存在被动技能、职业天赋、冲刺技能或额外功能键？',
    '5.3 小技能消耗什么资源？冷却、能量、魔力、弹药、充能次数，还是混合？',
    '5.4 大招如何释放或积攒？固定长 CD、造成伤害充能、击杀充能、受击充能，还是其他？',
    '5.5 闪避体力如何恢复？自动恢复、击杀恢复、装备词条恢复、休息事件恢复？',
    '5.6 玩家是否有生命、护盾、受击硬直、无敌保护、死亡复活？',
    '5.7 普通怪、精英怪、Boss 分别承担什么战斗体验？',
    '5.8 Boss 是否有阶段变化？每层 Boss 是否都是小 Boss，最终层是否是大 Boss？',
])
section_questions(doc, '6. 装备系统', [
    '6.1 每件装备需要包含哪些字段？例如 ID、名称、部位、品质、等级、基础属性、二级词条、技能词条、套装、SP 技能等。',
    '6.2 白、绿、蓝、紫、粉、橙、红各品质的词条数量、掉落概率、专属机制是否需要在完整 GDD 中展开？',
    '6.3 红装是否唯一？是否允许重复获得？不可删除、不可融合是否绝对成立？',
    '6.4 词条是否有品质、等级、类型、稀有度？是否可以被提取、镶嵌、锁定、洗练？',
    '6.5 提取词条是否消耗原装备？镶嵌是否有槽位限制？',
    '6.6 毕业装备的打造路径是否是“装备本体 + 高价值提取词条 + 洗练/镶嵌”？',
    '6.7 融合是局内系统还是局外系统？是否能指定目标装备或保留词条？',
    '6.8 洗练、提取、镶嵌、分解、出售、收藏各自的优先级是什么？',
])
section_questions(doc, '7. 技能与 Build 系统', [
    '7.1 技能来源的优先级如何排序？角色自带、武器、装备词条、套装、SP、转职、觉醒、技能书、局内事件。',
    '7.2 玩家最终携带几个主动技能？普攻是否算技能？闪避是否可被装备词条改造？',
    '7.3 技能裂变需要覆盖哪些维度？伤害、冷却、范围、数量、弹道、元素、分裂、穿透、追踪、爆炸、连锁、召唤物、持续区域、吸附控制、次级子弹等。',
    '7.4 转职是永久解锁后自由切换，还是每次选择后有成本？',
    '7.5 觉醒是角色维度还是职业维度？是否有多个分支？是否改变大招形态？',
    '7.6 技能书是局内临时效果，还是局外永久解锁后在家园配置？',
])
section_questions(doc, '8. 关卡、地图与路线', [
    '8.1 游戏预计有多少章节？每个章节是否有独立主题、怪物池、装备池和 Boss？',
    '8.2 每层目标房间数量是多少？第二层、第三层只是难度提升，还是事件池、怪物池、掉落池也变化？',
    '8.3 地图是固定列推进、网状节点，还是配置生成后以 DNF 房间方式体验？',
    '8.4 玩家是否允许回退、绕路、进入死路、钥匙门或锁定路线？',
    '8.5 迷雾预览能看到几步？能看到事件类型、奖励类型、危险等级、Boss 类型、商人类型吗？',
    '8.6 是否有道具、天赋或事件可以扩大预览范围？',
])
section_questions(doc, '9. 层结算、背包与带出', [
    '9.1 局内背包容量大概多大？装备和材料是否共用背包？',
    '9.2 背包满了是否还能拾取？是否支持自动拾取、筛选、排序、锁定？',
    '9.3 主动撤离是否等同成功带出？战斗中主动退出如何处理？',
    '9.4 出售是否只能在局内商店？层结算界面是否允许出售？',
    '9.5 带出格子提升只来自家园建筑，还是角色天赋、章节奖励、装备词条也可以影响？',
])
section_questions(doc, '10. 疲劳值系统', [
    '10.1 初始疲劳值和最大疲劳值大概是多少？最大疲劳是否可成长？',
    '10.2 进入下一层时疲劳是否恢复？击败 Boss 后恢复多少？',
    '10.3 0 疲劳是否仍可继续探索？如果可以，惩罚强度希望多高？',
    '10.4 疲劳影响哪些维度？攻击、移速、冷却、受伤、掉落、怪物强度、体力恢复、商店价格、事件概率？',
    '10.5 是否存在消耗疲劳换高收益的深渊房或高风险房？',
])
section_questions(doc, '11. 事件系统', [
    '11.1 第一批完整设计需要包含哪些事件？普通战斗、精英战斗、Boss、商人、休息、宝箱、深渊、高风险奖励、Boss 前准备、NPC、时装、角色、宠物？',
    '11.2 事件是否可以组合？例如商人+休息、深渊+宝箱、Boss 前商人+休息。',
    '11.3 事件是否受章节、层数、疲劳、难度、稀有度影响？',
    '11.4 事件是否能改变技能、Build 或局内临时强化？',
])
section_questions(doc, '12. 局外养成', [
    '12.1 家园核心功能有哪些？提升带出格子、制造疲劳药剂、装备加工、角色成长、宠物养成、时装展示、章节入口？',
    '12.2 家园是否有建筑升级、材料消耗、时间生产？',
    '12.3 角色是职业差异还是独立角色？是否有等级、天赋树、共享装备？',
    '12.4 宠物是战斗辅助、拾取辅助、属性加成，还是收集系统？',
    '12.5 时装是纯外观还是有属性/套装？是否通过事件获得？',
])
section_questions(doc, '13. 资源、经济与掉落', [
    '13.1 游戏资源有哪些？局内货币、局外金币、强化材料、融合材料、词条材料、洗练材料、疲劳药剂材料、时装材料、宠物材料、章节材料、觉醒/转职材料等。',
    '13.2 怪物、精英怪、Boss、事件、分解、出售分别产出什么？',
    '13.3 装备打造、家园、章节推进、技能、宠物、时装分别消耗什么？',
    '13.4 Boss 是否必掉高品质装备？是否有保底、章节掉落池、稀有掉落提示？',
])
section_questions(doc, '14. UI / UX、美术、音效', [
    '14.1 GDD 是否需要覆盖核心 UI 界面？主菜单、家园、章节、战斗 HUD、地图、背包、装备详情、层结算、死亡结算等。',
    '14.2 UI 风格偏暗黑、像素、清晰功能型，还是三者结合？',
    '14.3 装备对比、掉落品质提示、稀有掉落反馈是否需要重点设计？',
    '14.4 美术部分是否只写方向，还是需要拆到角色、怪物、Boss、装备、技能特效、场景、UI？',
    '14.5 音效与音乐是否需要进入完整 GDD？',
])
section_questions(doc, '15. 配置表与版本规划', [
    '15.1 是否希望 GDD 中明确所有需要配置化的数据表？装备、词条、套装、技能、怪物、Boss、掉落、关卡、事件、疲劳、家园、资源等。',
    '15.2 配置文件格式倾向 Excel、CSV、JSON，还是由程序决定？',
    '15.3 版本规划是否需要写入 GDD？例如原型验证、第一版 Demo、Alpha、Beta、正式版、后续扩展。',
    '15.4 完整 GDD 是否需要包含风险点清单和待定问题清单？',
])

p(doc, '')
doc.add_heading('三、推荐完整 GDD 框架', level=1)
table(doc, ['章节', '模块', '主要内容', '是否保留 / 修改意见'], [
    ['0', '文档信息', '修改记录、文档目的、适用范围、术语表', ''],
    ['1', '游戏总览', '产品定位、一句话介绍、目标平台、目标玩家、商业模式、核心参考、设计支柱', ''],
    ['2', '核心体验', '刷宝、Build、战斗、风险收益、长期养成体验目标', ''],
    ['3', '核心循环', '秒级循环、房间级循环、层级循环、单局循环、局外成长循环、结算循环', ''],
    ['4', '玩家系统', '属性、移动、普攻、技能槽、闪避、切目标、生命/护盾/体力/能量、受击死亡、输入方案', ''],
    ['5', '战斗系统', '伤害、弹幕、索敌、技能释放、冷却资源、Buff/Debuff、怪物、精英、Boss', ''],
    ['6', '技能与 Build 系统', '技能来源、技能槽、技能成长、技能裂变、SP、转职、觉醒、技能书、Build 成型路径', ''],
    ['7', '装备系统', '装备部位、品质、属性、词条、套装、SP、掉落、洗练、提取、镶嵌、融合、分解、出售', ''],
    ['8', '关卡与地图系统', '章节、层、房间、节点生成、路线选择、传送门、迷雾预览、房间类型、隐藏层/无尽层', ''],
    ['9', '层结算与带出系统', '层结算、带出、继续、出售、失败、背包、带出格子、风险收益曲线', ''],
    ['10', '疲劳值系统', '设计目的、消耗、恢复、阈值、增益、减益、路线选择、配置需求', ''],
    ['11', '事件系统', '商人、休息、深渊、宝箱、NPC、时装、角色、宠物、事件组合、扩展规则', ''],
    ['12', '怪物与 Boss 系统', '怪物目标、普通怪、精英怪、属性、技能、掉落、Boss 阶段、Boss 掉落、章节 Boss', ''],
    ['13', '局外养成系统', '家园、角色、职业/转职、觉醒、宠物、时装、章节推进、长期目标', ''],
    ['14', '资源与经济系统', '资源列表、局内/局外货币、装备材料、技能材料、家园材料、产出、消耗、经济闭环', ''],
    ['15', '掉落系统', '普通怪、精英怪、Boss、章节掉落池、品质概率、词条概率、保底、稀有掉落、表现', ''],
    ['16', 'UI / UX 设计', '主界面、家园、章节、战斗 HUD、地图、背包、装备详情、层结算、死亡结算、掉落反馈', ''],
    ['17', '美术与表现', '美术风格、角色、怪物、Boss、装备、技能特效、掉落特效、UI、场景氛围', ''],
    ['18', '音效与音乐', '战斗音效、技能音效、掉落音效、UI 音效、Boss 音乐、章节音乐', ''],
    ['19', '配置表设计', '装备表、词条表、套装表、技能表、怪物表、Boss 表、掉落表、关卡表、事件表、疲劳表等', ''],
    ['20', '版本规划', '原型验证、第一版 Demo、Alpha、Beta、正式版、后续扩展方向', ''],
    ['21', '附录', '名词解释、参考游戏拆解、待定问题清单、风险点清单、后续拆分文档列表', ''],
])
reply_box(doc, '对完整 GDD 框架的整体修改意见：')

p(doc, '')
doc.add_heading('四、GDD 系统关系图', level=1)
flow = '''游戏总览
  ↓
核心体验
  ↓
核心循环
  ├─ 战斗系统 → 技能与 Build 系统
  ├─ 装备系统 → 掉落系统 → 资源经济系统
  ├─ 关卡地图系统 → 事件系统 → 疲劳值系统
  ├─ 层结算与带出系统 → 局外养成系统
  └─ UI / UX、美术、音效、配置表支撑所有系统'''
p(doc, flow)
reply_box(doc, '对系统关系图的修改意见：')

p(doc, '')
doc.add_heading('五、玩家核心流程图', level=1)
player_flow = '''家园准备
  ↓
选择章节
  ↓
进入第 1 层
  ↓
选择路线节点（查看事件 / 奖励 / 危险）
  ↓
进入房间（消耗疲劳）
  ↓
战斗或事件处理
  ↓
拾取战利品
  ↓
是否到达 Boss？
  ├─ 否：继续选择路线节点
  └─ 是：进入 Boss 战
          ↓
       层结算
          ├─ 带出：返回家园养成
          ├─ 出售：获得局内货币后继续结算选择
          └─ 继续：进入下一层，难度提高、奖励提高、战利品继续承担风险'''
p(doc, player_flow)
reply_box(doc, '对玩家核心流程图的修改意见：')

p(doc, '')
doc.add_heading('六、建议推进顺序', level=1)
num(doc, [
    '确认完整 GDD 总框架。',
    '确认设计支柱和核心循环。',
    '确认战斗、装备、技能三大核心系统。',
    '确认关卡、层结算、疲劳、事件。',
    '确认局外养成、资源经济、掉落。',
    '确认 UI、美术、音效、配置表、版本规划。',
    '根据确认内容生成完整 GDD 初稿。',
])
reply_box(doc, '对推进顺序的修改意见：')

doc.save(out)
print(out.as_posix())
