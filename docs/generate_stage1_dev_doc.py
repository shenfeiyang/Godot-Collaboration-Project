from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

out = Path(r'c:/Users/loofnn/Desktop/sdc/策划案文档/第一版_demo说明文档/一阶段项目开发说明.docx')


def set_font(run, size=10.5, bold=False, color=None):
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def setup(doc):
    styles = doc.styles
    styles['Normal'].font.name = '宋体'
    styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    styles['Normal'].font.size = Pt(10.5)
    for name, size, color in [('Heading 1', 16, '1F4E79'), ('Heading 2', 13, '1F4E79'), ('Heading 3', 11, '000000')]:
        st = styles[name]
        st.font.name = '宋体'
        st._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
    for sec in doc.sections:
        sec.top_margin = Cm(2.2)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.4)
        sec.right_margin = Cm(2.4)


def p(doc, text='', style=None, size=10.5, bold=False, color=None, align=None):
    para = doc.add_paragraph(style=style)
    para.paragraph_format.line_spacing = 1.25
    if align:
        para.alignment = align
    r = para.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    return para


def bullet(doc, items):
    for item in items:
        para = doc.add_paragraph(style='List Bullet')
        para.paragraph_format.line_spacing = 1.18
        r = para.add_run(item)
        set_font(r)


def num(doc, items):
    for item in items:
        para = doc.add_paragraph(style='List Number')
        para.paragraph_format.line_spacing = 1.18
        r = para.add_run(item)
        set_font(r)


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run(h)
        set_font(r, bold=True)
        shade_cell(cell, 'D9EAF7')
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            para = cells[i].paragraphs[0]
            r = para.add_run(str(val))
            set_font(r, size=10)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return t


doc = Document()
setup(doc)

p(doc, '一阶段项目开发说明', size=22, bold=True, color='1F4E79', align=WD_ALIGN_PARAGRAPH.CENTER)
p(doc, '基于《SDC第一版Demo说明书》的开发细化版本', size=12, color='595959', align=WD_ALIGN_PARAGRAPH.CENTER)
p(doc, '')
table(doc, ['文档项', '内容'], [
    ['文档名称', '一阶段项目开发说明'],
    ['所属阶段', '第一版 Demo / 一阶段开发'],
    ['文档定位', '将 Demo 目标细化为可落地的阶段开发说明'],
    ['核心目标', '验证关卡内循环'],
    ['不接入内容', '复杂外围养成、完整家园、宠物正式系统、时装正式系统'],
])

doc.add_heading('1. 一阶段开发目标', level=1)
p(doc, '一阶段开发的目标不是完成 SDC 的完整产品，而是围绕“关卡内循环”做出可验证、可反复游玩的第一版 Demo。')
bullet(doc, [
    '验证玩家从进入章节到通关或失败结算的一整轮关卡体验。',
    '验证房间推进、战斗、掉落、Build 成型、Boss、翻牌、商店、传输与结算的闭环。',
    '验证装备驱动技能变化是否足以构成“继续刷”的动力。',
    '验证自选难度 Tag 是否能提供清晰的风险收益差异。',
])

doc.add_heading('2. 一阶段开发范围', level=1)
table(doc, ['模块', '开发范围', '备注'], [
    ['章节', '1 个章节', '用于验证单章节完整循环。'],
    ['层数', '3 层', '每层存在关底 Boss。'],
    ['地图推进', '哈迪斯-like 三选一房间推进', '路线不可返回。'],
    ['房间类型', '普通战斗、精英战斗、Boss、商人、休息、深渊', '可按开发进度裁剪休息/深渊表现复杂度。'],
    ['角色', '优先 1 个基础职业路线', '建议枪手。'],
    ['技能', '普攻 + 3 小技能 + 1 觉醒', '满足基础 Build 体验。'],
    ['怪物', '普通怪、精英怪、每层 Boss', '深渊可共用怪物库强化。'],
    ['装备', '基础品质、基础掉落、基础词条、基础技能裂变', '不做完整终局词条体系。'],
    ['结算', '翻牌奖励、道中商店、对外传输、全局胜败带出', '必须接入。'],
    ['难度系统', '基础自选难度 Tag', '至少支持若干简单 Tag。'],
])

doc.add_heading('3. 一阶段不开发或弱化的内容', level=1)
bullet(doc, [
    '复杂家园系统。',
    '宠物正式系统。',
    '时装正式系统。',
    '多章节内容。',
    '复杂长期资源循环。',
    '复杂职业矩阵与大量转职路线。',
    '完整装备终局系统与复杂保护道具经济。',
])

doc.add_heading('4. 一阶段核心玩法闭环', level=1)
num(doc, [
    '玩家选择章节。',
    '玩家选择若干自选难度 Tag。',
    '玩家进入第 1 层。',
    '玩家完成当前房间战斗或事件目标。',
    '玩家手动拾取掉落。',
    '玩家在三选一传送门中选择下一个房间。',
    '玩家推进至关底 Boss。',
    '玩家击败 Boss 后获得翻牌奖励。',
    '玩家进入道中商店，出售或购买局内道具。',
    '玩家执行一次道具对外传输。',
    '玩家进入下一层，直至通关、失败或主动退出。',
    '玩家根据胜利 / 失败结算带出部分奖励。',
])

doc.add_heading('5. 必须完成的系统模块（P0）', level=1)
table(doc, ['P0 模块', '说明'], [
    ['关卡推进框架', '实现章节进入、房间完成、三选一传送门、Boss 房推进。'],
    ['玩家基础战斗', '移动、普攻、3 小技能、1 觉醒、闪避、索敌、受击、死亡。'],
    ['怪物基础行为', '普通怪、精英怪、Boss 的基础 AI 与战斗表现。'],
    ['掉落与拾取', '怪物掉落、手动拾取、掉落进背包。'],
    ['局内背包', '共用背包、容量限制、基础筛选/排序/锁定。'],
    ['基础装备结构', '装备品质、穿戴、基础词条、普攻/技能增益读取。'],
    ['基础 Build 裂变', '装备对技能和普攻产生可感知变化。'],
    ['Boss 奖励结算', '翻牌奖励、道中商店、对外传输。'],
    ['全局胜败结算', '主动退出、失败、通关三类结果都能成立。'],
])

doc.add_heading('6. 应尽量完成的系统模块（P1）', level=1)
table(doc, ['P1 模块', '说明'], [
    ['自选难度 Tag', '进入章节前选择若干 Tag，影响掉落与事件权重。'],
    ['基础事件扩展', '商人、休息、深渊事件形成基础差异。'],
    ['特殊传送门', '提供额外特殊房间入口。'],
    ['掉落反馈 UI', '稀有装备掉落、奖励提示、房间信息展示。'],
    ['更多装备裂变样例', '至少形成若干清晰可感知的 Build 方向。'],
])

doc.add_heading('7. 可后置的模块（P2）', level=1)
bullet(doc, [
    '基础家园接口。',
    '第二基础职业。',
    '更丰富的深渊与隐藏内容。',
    '宠物辅助拾取雏形。',
    '时装外观占位。',
])

doc.add_heading('8. 推荐开发顺序', level=1)
num(doc, [
    '完成基础地图推进与房间切换。',
    '完成玩家移动、普攻、技能、怪物战斗。',
    '完成基础掉落、拾取、背包。',
    '完成基础装备系统与技能裂变读取。',
    '完成每层 Boss 与胜败判定。',
    '完成翻牌奖励、道中商店、对外传输。',
    '完成自选难度 Tag。',
    '补齐深渊、休息、特殊传送门等扩展事件。',
])

doc.add_heading('9. 角色与战斗实现建议', level=1)
bullet(doc, [
    '第一版优先只做 1 个职业路线，建议枪手。',
    '优先做“可读性强、差异明显”的普攻和技能，便于验证装备裂变价值。',
    '第一版 Build 重点不是数量，而是清晰度。至少保证玩家能明显感知到 2~3 种不同战斗变化。',
    '战斗爽感以输出为先，防御只作为高难挑战容错。',
])

doc.add_heading('10. 房间与关卡实现建议', level=1)
table(doc, ['模块', '建议'], [
    ['房间数', '单层控制在约 8~10 个房间。'],
    ['层定位', '第 1 层成型，第 2 层强化，第 3 层检验。'],
    ['房间展示', '向玩家展示房间类型、奖励类型、危险等级、是否精英/深渊/特殊事件。'],
    ['事件取舍', '若工期有限，可先保证普通战斗、精英、Boss、商人四类完整。'],
    ['特殊传送门', '可作为一阶段加分项，不是最先必须完成。'],
])

doc.add_heading('11. 装备与掉落实现建议', level=1)
bullet(doc, [
    '第一版不需要完整终局装备体系，但必须让掉落有明显层次。',
    '优先保证：掉落 → 穿戴 → 战斗变化 → 更想继续刷 的最短闭环。',
    '基础词条建议优先覆盖：伤害、冷却、技能数量、分裂、穿透、追踪、范围、元素变化。',
    '稀有装备和高价值词条需要有明显的 UI 反馈。',
])

doc.add_heading('12. 一阶段成功标准', level=1)
bullet(doc, [
    '玩家能完整体验 1 个章节 3 层流程。',
    '玩家能清晰理解三选一路线推进。',
    '玩家能感知到装备改变技能或普攻表现。',
    '玩家能理解翻牌奖励、道中商店、对外传输和胜败带出的价值。',
    '玩家能感受到自选难度 Tag 带来的收益差异。',
    '这一轮 Demo 足以证明关卡内循环具有继续扩展的价值。',
])

doc.add_heading('13. 一阶段风险点', level=1)
bullet(doc, [
    '若装备裂变不明显，玩家无法形成刷宝动力。',
    '若房间推进过慢，3 层流程会显得冗长。',
    '若翻牌、商店、传输流程不清晰，玩家难以理解结算价值。',
    '若自选难度 Tag 与奖励反馈脱节，风险收益将不成立。',
    '若 Boss 强度和 Build 成长曲线不匹配，会削弱关卡节奏。',
])

doc.add_heading('14. 与后续文档的衔接', level=1)
p(doc, '《一阶段项目开发说明》只负责第一版 Demo 的开发范围和落地顺序。具体战斗公式、技能表、装备表、怪物表、关卡表等内容，仍需要在后续专项策划案中展开。')
table(doc, ['后续文档', '说明'], [
    ['战斗规则说明', '补充战斗公式、属性和表现规则。'],
    ['技能说明', '细化具体技能与裂变方式。'],
    ['装备系统策划案', '细化装备字段、品质、词条与装备成长。'],
    ['关卡策划案', '细化章节、房间库、事件权重、特殊传送门。'],
    ['怪物策划案', '细化普通怪、精英怪、Boss 行为与掉落。'],
])

doc.save(out)
print(out.as_posix())
